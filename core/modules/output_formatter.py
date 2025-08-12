"""
Multi-Format Output System for Jarvis.
Exports results to PDF, Word, Excel, and other formats.
"""
import json
import csv
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
from core.utils.log import logger

# PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel generation
try:
    import pandas as pd
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils.dataframe import dataframe_to_rows
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# HTML to PDF
try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False


class OutputFormatter:
    def __init__(self):
        self.output_dir = Path("exports")
        self.output_dir.mkdir(exist_ok=True)
        
        # Check available libraries
        self.available_formats = {
            'json': True,
            'csv': True,
            'txt': True,
            'html': True,
            'pdf': REPORTLAB_AVAILABLE or PDFKIT_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'xlsx': EXCEL_AVAILABLE
        }
        
        logger.info(f"Available export formats: {[k for k, v in self.available_formats.items() if v]}")
    
    def export_data(self, data: Dict[str, Any], filename: str, formats: List[str] = None) -> Dict[str, str]:
        """Export data to multiple formats."""
        try:
            if not formats:
                formats = ['json', 'pdf', 'xlsx']  # Default formats
            
            exported_files = {}
            base_filename = Path(filename).stem
            
            for format_type in formats:
                if not self.available_formats.get(format_type, False):
                    logger.warning(f"Format {format_type} not available")
                    continue
                
                try:
                    if format_type == 'json':
                        file_path = self._export_json(data, base_filename)
                    elif format_type == 'csv':
                        file_path = self._export_csv(data, base_filename)
                    elif format_type == 'txt':
                        file_path = self._export_txt(data, base_filename)
                    elif format_type == 'html':
                        file_path = self._export_html(data, base_filename)
                    elif format_type == 'pdf':
                        file_path = self._export_pdf(data, base_filename)
                    elif format_type == 'docx':
                        file_path = self._export_docx(data, base_filename)
                    elif format_type == 'xlsx':
                        file_path = self._export_xlsx(data, base_filename)
                    else:
                        continue
                    
                    exported_files[format_type] = str(file_path)
                    logger.info(f"✅ Exported {format_type}: {file_path}")
                    
                except Exception as e:
                    logger.error(f"Failed to export {format_type}: {e}")
                    continue
            
            return {
                'success': True,
                'exported_files': exported_files,
                'total_formats': len(exported_files)
            }
            
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def _export_json(self, data: Dict, filename: str) -> Path:
        """Export to JSON format."""
        file_path = self.output_dir / f"{filename}.json"
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, default=str, ensure_ascii=False)
        
        return file_path
    
    def _export_csv(self, data: Dict, filename: str) -> Path:
        """Export to CSV format."""
        file_path = self.output_dir / f"{filename}.csv"
        
        # Flatten data for CSV export
        flattened_data = self._flatten_data_for_csv(data)
        
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            if flattened_data:
                writer = csv.DictWriter(f, fieldnames=flattened_data[0].keys())
                writer.writeheader()
                writer.writerows(flattened_data)
        
        return file_path
    
    def _export_txt(self, data: Dict, filename: str) -> Path:
        """Export to plain text format."""
        file_path = self.output_dir / f"{filename}.txt"
        
        text_content = self._convert_to_text(data)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        return file_path
    
    def _export_html(self, data: Dict, filename: str) -> Path:
        """Export to HTML format."""
        file_path = self.output_dir / f"{filename}.html"
        
        html_content = self._convert_to_html(data)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return file_path
    
    def _export_pdf(self, data: Dict, filename: str) -> Path:
        """Export to PDF format."""
        file_path = self.output_dir / f"{filename}.pdf"
        
        if REPORTLAB_AVAILABLE:
            return self._export_pdf_reportlab(data, file_path)
        elif PDFKIT_AVAILABLE:
            return self._export_pdf_pdfkit(data, file_path)
        else:
            raise ImportError("No PDF library available")
    
    def _export_pdf_reportlab(self, data: Dict, file_path: Path) -> Path:
        """Export to PDF using ReportLab."""
        doc = SimpleDocTemplate(str(file_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title = data.get('title', 'Jarvis Export Report')
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        if 'metadata' in data:
            story.append(Paragraph("Report Information", styles['Heading2']))
            metadata = data['metadata']
            for key, value in metadata.items():
                story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Main content
        story.append(Paragraph("Report Content", styles['Heading2']))
        
        # Convert data to readable format
        content_text = self._convert_to_text(data.get('data', data))
        
        # Split into paragraphs and add to story
        paragraphs = content_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        return file_path
    
    def _export_docx(self, data: Dict, filename: str) -> Path:
        """Export to Word document format."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        file_path = self.output_dir / f"{filename}.docx"
        
        doc = Document()
        
        # Title
        title = data.get('title', 'Jarvis Export Report')
        doc.add_heading(title, 0)
        
        # Metadata
        if 'metadata' in data:
            doc.add_heading('Report Information', level=1)
            metadata = data['metadata']
            for key, value in metadata.items():
                p = doc.add_paragraph()
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                p.add_run(str(value))
        
        # Main content
        doc.add_heading('Report Content', level=1)
        
        # Add content based on data structure
        self._add_content_to_docx(doc, data.get('data', data))
        
        doc.save(str(file_path))
        return file_path
    
    def _export_xlsx(self, data: Dict, filename: str) -> Path:
        """Export to Excel format."""
        if not EXCEL_AVAILABLE:
            raise ImportError("pandas and openpyxl not available")
        
        file_path = self.output_dir / f"{filename}.xlsx"
        
        with pd.ExcelWriter(str(file_path), engine='openpyxl') as writer:
            # Metadata sheet
            if 'metadata' in data:
                metadata_df = pd.DataFrame(list(data['metadata'].items()), columns=['Property', 'Value'])
                metadata_df.to_excel(writer, sheet_name='Metadata', index=False)
            
            # Data sheets
            data_content = data.get('data', {})
            
            if isinstance(data_content, dict):
                for sheet_name, sheet_data in data_content.items():
                    try:
                        if isinstance(sheet_data, list) and sheet_data:
                            df = pd.DataFrame(sheet_data)
                            # Clean sheet name for Excel
                            clean_sheet_name = self._clean_sheet_name(sheet_name)
                            df.to_excel(writer, sheet_name=clean_sheet_name, index=False)
                    except Exception as e:
                        logger.warning(f"Could not export sheet {sheet_name}: {e}")
            
            # Summary sheet
            summary_data = self._create_summary_data(data)
            if summary_data:
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='Summary', index=False)
        
        return file_path
    
    def _flatten_data_for_csv(self, data: Dict) -> List[Dict]:
        """Flatten nested data structure for CSV export."""
        flattened = []
        
        def flatten_item(item, prefix=''):
            if isinstance(item, dict):
                result = {}
                for key, value in item.items():
                    new_key = f"{prefix}_{key}" if prefix else key
                    if isinstance(value, (dict, list)):
                        result.update(flatten_item(value, new_key))
                    else:
                        result[new_key] = value
                return result
            elif isinstance(item, list):
                result = {}
                for i, value in enumerate(item):
                    new_key = f"{prefix}_{i}" if prefix else str(i)
                    if isinstance(value, (dict, list)):
                        result.update(flatten_item(value, new_key))
                    else:
                        result[new_key] = value
                return result
            else:
                return {prefix: item} if prefix else {'value': item}
        
        # Handle different data structures
        if 'data' in data and isinstance(data['data'], dict):
            for category, items in data['data'].items():
                if isinstance(items, list):
                    for item in items:
                        flat_item = flatten_item(item)
                        flat_item['category'] = category
                        flattened.append(flat_item)
                else:
                    flat_item = flatten_item(items)
                    flat_item['category'] = category
                    flattened.append(flat_item)
        else:
            flattened.append(flatten_item(data))
        
        return flattened
    
    def _convert_to_text(self, data: Any, indent: int = 0) -> str:
        """Convert data structure to readable text."""
        text_parts = []
        indent_str = "  " * indent
        
        if isinstance(data, dict):
            for key, value in data.items():
                if key in ['soup', 'response']:  # Skip non-serializable objects
                    continue
                
                text_parts.append(f"{indent_str}{key.replace('_', ' ').title()}:")
                
                if isinstance(value, (dict, list)):
                    text_parts.append(self._convert_to_text(value, indent + 1))
                else:
                    text_parts.append(f"{indent_str}  {value}")
                text_parts.append("")
        
        elif isinstance(data, list):
            for i, item in enumerate(data):
                text_parts.append(f"{indent_str}Item {i + 1}:")
                text_parts.append(self._convert_to_text(item, indent + 1))
                text_parts.append("")
        
        else:
            text_parts.append(f"{indent_str}{data}")
        
        return "\n".join(text_parts)
    
    def _convert_to_html(self, data: Dict) -> str:
        """Convert data to HTML format."""
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            "<meta charset='utf-8'>",
            f"<title>{data.get('title', 'Jarvis Export Report')}</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; margin: 40px; }",
            "h1 { color: #333; border-bottom: 2px solid #007acc; }",
            "h2 { color: #555; margin-top: 30px; }",
            "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #f2f2f2; }",
            ".metadata { background-color: #f9f9f9; padding: 15px; border-radius: 5px; }",
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{data.get('title', 'Jarvis Export Report')}</h1>"
        ]
        
        # Add metadata
        if 'metadata' in data:
            html_parts.append("<div class='metadata'>")
            html_parts.append("<h2>Report Information</h2>")
            for key, value in data['metadata'].items():
                html_parts.append(f"<p><strong>{key.replace('_', ' ').title()}:</strong> {value}</p>")
            html_parts.append("</div>")
        
        # Add main content
        html_parts.append("<h2>Report Content</h2>")
        html_parts.append(self._data_to_html(data.get('data', data)))
        
        html_parts.extend(["</body>", "</html>"])
        
        return "\n".join(html_parts)
    
    def _data_to_html(self, data: Any) -> str:
        """Convert data to HTML representation."""
        if isinstance(data, dict):
            html = "<div>"
            for key, value in data.items():
                if key in ['soup', 'response']:  # Skip non-serializable objects
                    continue
                html += f"<h3>{key.replace('_', ' ').title()}</h3>"
                html += self._data_to_html(value)
            html += "</div>"
            return html
        
        elif isinstance(data, list) and data:
            if isinstance(data[0], dict):
                # Create table for list of dictionaries
                html = "<table>"
                # Header
                html += "<tr>"
                for key in data[0].keys():
                    html += f"<th>{key.replace('_', ' ').title()}</th>"
                html += "</tr>"
                # Rows
                for item in data:
                    html += "<tr>"
                    for value in item.values():
                        html += f"<td>{value}</td>"
                    html += "</tr>"
                html += "</table>"
                return html
            else:
                # Simple list
                html = "<ul>"
                for item in data:
                    html += f"<li>{item}</li>"
                html += "</ul>"
                return html
        
        else:
            return f"<p>{data}</p>"
    
    def _add_content_to_docx(self, doc, data: Any):
        """Add content to Word document."""
        if isinstance(data, dict):
            for key, value in data.items():
                if key in ['soup', 'response']:  # Skip non-serializable objects
                    continue
                
                doc.add_heading(key.replace('_', ' ').title(), level=2)
                self._add_content_to_docx(doc, value)
        
        elif isinstance(data, list) and data:
            if isinstance(data[0], dict):
                # Create table for list of dictionaries
                table = doc.add_table(rows=1, cols=len(data[0].keys()))
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(data[0].keys()):
                    hdr_cells[i].text = key.replace('_', ' ').title()
                
                # Data rows
                for item in data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(item.values()):
                        row_cells[i].text = str(value)
            else:
                # Simple list
                for item in data:
                    doc.add_paragraph(str(item), style='List Bullet')
        
        else:
            doc.add_paragraph(str(data))
    
    def _create_summary_data(self, data: Dict) -> List[Dict]:
        """Create summary data for Excel export."""
        summary = []
        
        # Basic statistics
        if 'data' in data:
            for category, items in data['data'].items():
                if isinstance(items, list):
                    summary.append({
                        'Category': category.replace('_', ' ').title(),
                        'Count': len(items),
                        'Type': 'List'
                    })
                else:
                    summary.append({
                        'Category': category.replace('_', ' ').title(),
                        'Count': 1,
                        'Type': 'Object'
                    })
        
        return summary
    
    def _clean_sheet_name(self, name: str) -> str:
        """Clean sheet name for Excel compatibility."""
        # Excel sheet names can't contain certain characters
        invalid_chars = ['\\', '/', '*', '[', ']', ':', '?']
        clean_name = name
        for char in invalid_chars:
            clean_name = clean_name.replace(char, '_')
        
        # Limit to 31 characters (Excel limit)
        return clean_name[:31]
    
    def get_available_formats(self) -> List[str]:
        """Get list of available export formats."""
        return [format_type for format_type, available in self.available_formats.items() if available]
